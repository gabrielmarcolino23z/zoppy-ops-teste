import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
import markdown
from bs4 import BeautifulSoup
import io
import re

def setup_document_styles(doc):
    """
    Setup document styles including default font and heading styles
    """
    # Configuração da fonte padrão (Arial)
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)  # Tamanho padrão para texto normal
    style.font.color.rgb = RGBColor(0, 0, 0)  # Preto
    
    # Heading 1 (#)
    h1_style = doc.styles['Heading 1']
    h1_style.font.size = Pt(20)
    h1_style.font.bold = True
    h1_style.font.name = 'Arial'
    h1_style.font.color.rgb = RGBColor(0, 0, 0)  # Preto
    
    # Heading 2 (##)
    h2_style = doc.styles['Heading 2']
    h2_style.font.size = Pt(16)  
    h2_style.font.bold = True
    h2_style.font.name = 'Arial'
    h2_style.font.color.rgb = RGBColor(0, 0, 0)  # Preto
    
    # Heading 3 (###)
    h3_style = doc.styles['Heading 3']
    h3_style.font.size = Pt(14)  
    h3_style.font.bold = True
    h3_style.font.name = 'Arial'
    h3_style.font.color.rgb = RGBColor(0, 0, 0) 

def add_formatted_text(paragraph, text):
    """
    Adiciona texto ao parágrafo com formatação de negrito para textos entre **
    """
    # Remove espaços extras no início e fim e remove caracteres '>'
    text = text.strip().replace('>', '').strip()
    
    # Divide o texto em partes, separando pelos marcadores de negrito
    parts = re.split(r'(\*\*)', text)
    
    # Flag para controlar se estamos em uma parte em negrito
    is_bold = False
    
    # Processa cada parte
    for part in parts:
        if part:  # Ignora partes vazias
            if part == '**':
                # Encontramos um marcador de negrito, então alternamos o flag
                is_bold = not is_bold
            else:
                # Remove qualquer '>' remanescente e adiciona o texto com a formatação atual
                cleaned_part = part.replace('>', '').strip()
                run = paragraph.add_run(cleaned_part)
                run.font.name = 'Arial'
                # Só define o tamanho 11pt se não for um título
                if not paragraph.style.name.startswith('Heading'):
                    run.font.size = Pt(11)
                run.bold = is_bold

def set_cell_border(cell, **kwargs):
    """
    Set cell border
    Usage:
        set_cell_border(
            cell,
            top={"sz": 12, "color": "#000000", "val": "single"},
            bottom={"sz": 12, "color": "#000000", "val": "single"},
            start={"sz": 12, "color": "#000000", "val": "single"},
            end={"sz": 12, "color": "#000000", "val": "single"}
        )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # check for tag existance, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    
    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            
            # check for tag existance, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            
            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def get_column_widths(table, headers, cells):
    """
    Calcula as larguras das colunas de forma inteligente baseado no tipo e tamanho do conteúdo
    """
    # Largura total disponível (A4 - margens)
    total_width = 6.0  # Mantemos 6 polegadas como largura ideal da tabela
    
    # Analisa o conteúdo de cada coluna
    col_contents = []
    max_lengths = []
    content_types = []  # Tipo de conteúdo predominante em cada coluna
    
    for i in range(len(headers)):
        # Coleta todo o conteúdo da coluna
        col_content = [headers[i] if isinstance(headers[i], str) else headers[i].get_text().strip()]
        for row in cells:
            if i < len(row):
                cell_text = row[i] if isinstance(row[i], str) else row[i].get_text().strip()
                col_content.append(cell_text)
        col_contents.append(col_content)
        
        # Calcula o comprimento máximo do texto
        max_length = max(len(text) for text in col_content)
        max_lengths.append(max_length)
        
        # Analisa o tipo de conteúdo
        content_type = {
            'numeric': 0,
            'short_text': 0,  # < 15 caracteres
            'medium_text': 0,  # 15-30 caracteres
            'long_text': 0     # > 30 caracteres
        }
        
        for text in col_content:
            if text.replace('%', '').replace('-', '').strip().isdigit():
                content_type['numeric'] += 1
            elif len(text) < 15:
                content_type['short_text'] += 1
            elif len(text) < 30:
                content_type['medium_text'] += 1
            else:
                content_type['long_text'] += 1
        
        # Determina o tipo predominante
        predominant_type = max(content_type.items(), key=lambda x: x[1])[0]
        content_types.append(predominant_type)
    
    # Calcula os pesos base para cada coluna
    weights = []
    for i, content_type in enumerate(content_types):
        if content_type == 'numeric':
            base_weight = 0.5 
        elif content_type == 'short_text':
            base_weight = 0.8
        elif content_type == 'medium_text':
            base_weight = 1.2
        else:  # long_text
            base_weight = 1.5
            
        # Ajusta o peso baseado no comprimento máximo relativo
        length_factor = max_lengths[i] / max(max_lengths)
        weight = base_weight * (0.5 + 0.5 * length_factor)  # Combina tipo e comprimento
        weights.append(weight)
    
    # Normaliza os pesos para distribuir a largura total
    total_weight = sum(weights)
    widths = []
    
    for weight in weights:
        # Calcula a largura proporcional
        width = (weight / total_weight) * total_width
        
        # Aplica limites mínimos e máximos
        width = max(min(width, 3.0), 0.6)  # Mínimo de 0.6", máximo de 3.0"
        widths.append(width)
    
    # Ajusta as larguras para garantir que somem exatamente a largura total
    current_total = sum(widths)
    if current_total != total_width:
        # Distribui a diferença proporcionalmente
        factor = total_width / current_total
        widths = [w * factor for w in widths]
    
    return [Inches(w) for w in widths]

def format_table(table):
    """
    Apply consistent formatting to a table
    """
    # Define border styles
    border = {
        "sz": 4,  # Aumentei a espessura da borda
        "val": "single",
        "color": "#000000",
    }
    
    # Centraliza a tabela na página
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Aplica estilo à tabela inteira
    table.style = 'Table Grid'
    
    # Apply to each cell
    for row_idx, row in enumerate(table.rows):
        # Define altura da linha
        row.height = Inches(0.4)  # Altura padrão da linha
        
        for cell in row.cells:
            # Apply borders to each cell
            set_cell_border(
                cell,
                top=border,
                bottom=border,
                start=border,
                end=border,
            )
            
            if row_idx == 0:  # Primeira linha é sempre o cabeçalho
                cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="FFFFFF"/>'.format(nsdecls('w'))))
            
            # Set paragraph formatting
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Aumenta o espaçamento vertical
            paragraph.paragraph_format.space_before = Pt(8)
            paragraph.paragraph_format.space_after = Pt(8)
            
            # Adiciona espaçamento entre linhas
            paragraph.paragraph_format.line_spacing = Pt(14)
            
            # Set font
            for run in paragraph.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(11)  # Tamanho 11pt apenas para texto em tabelas
                if row_idx == 0:
                    run.font.bold = True

def convert_markdown_to_docx(markdown_content: str, company_name: str) -> bytes:
    """
    Convert markdown content to a DOCX document with minimal styling
    """
    doc = Document()
    setup_document_styles(doc)
    
    # Set up page size and margins (A4)
    section = doc.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    
    # Pré-processamento do markdown antes de converter para HTML
    lines = markdown_content.split('\n')
    processed_lines = []
    
    # Padrões para identificação de listas
    numbered_list_pattern = re.compile(r'^(\d+)\.\s+(.+)$')
    bullet_list_pattern = re.compile(r'^[-*]\s+(.+)$')
    in_list = False
    list_items = []
    list_level = 0
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Remove caracteres '>' de todas as linhas
        line = line.replace('>', '').strip()
        
        # Verifica se é uma linha horizontal (---)
        if line == '---':
            processed_lines.append("<!-- HORIZONTAL_RULE -->")
            i += 1
            continue
        
        # Verifica se é cabeçalho
        if line.startswith('# ') or line.startswith('## ') or line.startswith('### '):
            if in_list and list_items:
                processed_lines.append("<!-- LIST_START -->")
                processed_lines.extend(list_items)
                processed_lines.append("<!-- LIST_END -->")
                in_list = False
                list_items = []
            processed_lines.append(line)
        
        # Verifica se é item de lista numerada
        elif numbered_list_pattern.match(line):
            match = numbered_list_pattern.match(line)
            if not in_list:
                in_list = True
                list_items = []
            
            number = match.group(1)
            content = match.group(2).strip().replace('>', '')  # Remove '>' e espaços extras
            # Adiciona marcadores de negrito ao conteúdo
            if ':' in content:  # Se for um título (contém ':')
                title_part = content.split(':')[0] + ':'  # Pega a parte antes do ':' e adiciona o ':'
                rest_part = content.split(':', 1)[1] if len(content.split(':', 1)) > 1 else ''  # Resto do conteúdo
                content = f"**{title_part}**{rest_part}"
            list_items.append(f"<!-- NUMBERED_ITEM {number} -->{content}")
            
            # Verifica subitens
            j = i + 1
            while j < len(lines) and (lines[j].strip().startswith('-') or lines[j].strip().startswith('*') or lines[j].strip() == '' or lines[j].startswith('   ')):
                subline = lines[j].strip()
                if subline:  # Ignora linhas vazias
                    # Remove caracteres '>' dos subitens
                    subline = subline.replace('>', '').strip()
                    if subline.startswith('-') or subline.startswith('*'):
                        bullet_match = bullet_list_pattern.match(subline)
                        if bullet_match:
                            bullet_content = bullet_match.group(1).strip().replace('>', '')
                            list_items.append(f"<!-- BULLET_ITEM -->{bullet_content}")
                    else:
                        list_items.append(f"<!-- CONTINUATION -->{subline}")
                j += 1
            i = j - 1
        
        # Verifica se é um bullet point
        elif bullet_list_pattern.match(line):
            if in_list and not list_items[-1].startswith("<!-- BULLET_ITEM -->"):
                processed_lines.append("<!-- LIST_START -->")
                processed_lines.extend(list_items)
                processed_lines.append("<!-- LIST_END -->")
                in_list = False
                list_items = []
            
            match = bullet_list_pattern.match(line)
            if not in_list:
                in_list = True
                list_items = []
            
            content = match.group(1).strip().replace('>', '')  # Remove '>' e espaços extras
            list_items.append(f"<!-- BULLET_ITEM -->{content}")
        
        else:
            if in_list and list_items:
                processed_lines.append("<!-- LIST_START -->")
                processed_lines.extend(list_items)
                processed_lines.append("<!-- LIST_END -->")
                in_list = False
                list_items = []
            processed_lines.append(line)
        
        i += 1
    
    if in_list and list_items:
        processed_lines.append("<!-- LIST_START -->")
        processed_lines.extend(list_items)
        processed_lines.append("<!-- LIST_END -->")
    
    processed_markdown = '\n'.join(processed_lines)
    
    html = markdown.markdown(
        processed_markdown,
        extensions=['tables', 'nl2br', 'fenced_code']
    )
    
    soup = BeautifulSoup(html, 'html.parser')
    
    i = 0
    current_list_number = 1
    
    while i < len(processed_lines):
        line = processed_lines[i]
        
        if line == "<!-- HORIZONTAL_RULE -->":
            # Adiciona uma linha horizontal estilizada
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(20)  # Espaço antes da linha
            paragraph.paragraph_format.space_after = Pt(20)   # Espaço depois da linha
            
            # Adiciona uma borda na parte inferior do parágrafo
            pPr = paragraph._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '8')  # 8 pontos de espessura
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), '000000')
            pBdr.append(bottom)
            pPr.append(pBdr)
            
            # Adiciona um espaço após a linha
            doc.add_paragraph().paragraph_format.space_after = Pt(8)
            i += 1
            continue
        
        if line.startswith('# '):
            heading = doc.add_heading('', level=1)
            add_formatted_text(heading, f"**{line[2:].strip()}**")
            heading.paragraph_format.space_before = Pt(14)
            heading.paragraph_format.space_after = Pt(8)
        
        elif line.startswith('## '):
            heading = doc.add_heading('', level=2)
            add_formatted_text(heading, f"**{line[3:].strip()}**")
            heading.paragraph_format.space_before = Pt(14)
            heading.paragraph_format.space_after = Pt(8)
        
        elif line.startswith('### '):
            heading = doc.add_heading('', level=3)
            add_formatted_text(heading, f"**{line[4:].strip()}**")
            heading.paragraph_format.space_before = Pt(14)
            heading.paragraph_format.space_after = Pt(8)
        
        elif line == "<!-- LIST_START -->":
            list_items = []
            i += 1
            
            while i < len(processed_lines) and processed_lines[i] != "<!-- LIST_END -->":
                list_items.append(processed_lines[i])
                i += 1
            
            for item in list_items:
                if item.startswith("<!-- NUMBERED_ITEM "):
                    match = re.match(r'<!-- NUMBERED_ITEM (\d+) -->(.+)', item)
                    if match:
                        number = match.group(1)
                        content = match.group(2).strip()  # Remove espaços extras
                        p = doc.add_paragraph()
                        p.paragraph_format.left_indent = Inches(0)
                        add_formatted_text(p, f"{number}. {content}")
                        p.paragraph_format.space_after = Pt(4)
                
                elif item.startswith("<!-- BULLET_ITEM -->"):
                    content = item[19:].strip()  # Remove espaços extras
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.5)  # Aumentado de 0.25 para 0.5
                    add_formatted_text(p, f"• {content}")  # Usa bullet point unicode
                    p.paragraph_format.space_after = Pt(4)
                
                elif item.startswith("<!-- CONTINUATION -->"):
                    content = item[21:].strip()  # Remove espaços extras
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.5)  # Aumentado de 0.25 para 0.5
                    add_formatted_text(p, content)
                    p.paragraph_format.space_after = Pt(4)
        
        elif line.startswith('|') and '|' in line[1:]:
            table_lines = [line]
            j = i + 1
            while j < len(processed_lines) and processed_lines[j].startswith('|'):
                table_lines.append(processed_lines[j])
                j += 1
            
            table_markdown = '\n'.join(table_lines)
            table_html = markdown.markdown(table_markdown, extensions=['tables'])
            table_soup = BeautifulSoup(table_html, 'html.parser')
            table_element = table_soup.find('table')
            
            if table_element:
                headers = []
                header_row = table_element.find('tr')
                if header_row:
                    for th in header_row.find_all(['th', 'td']):
                        headers.append(th.get_text().strip())
                
                rows = []
                for tr in table_element.find_all('tr')[1:]:
                    cells = []
                    for td in tr.find_all('td'):
                        cells.append(td.get_text().strip())
                    if cells:
                        rows.append(cells)
                
                if headers and rows:
                    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
                    
                    for k, header in enumerate(headers):
                        cell = table.cell(0, k)
                        paragraph = cell.paragraphs[0]
                        add_formatted_text(paragraph, header)
                    
                    for row_idx, row in enumerate(rows, start=1):
                        for col_idx, cell_text in enumerate(row):
                            if col_idx < len(headers):
                                cell = table.cell(row_idx, col_idx)
                                paragraph = cell.paragraphs[0]
                                add_formatted_text(paragraph, cell_text)
                    
                    column_widths = get_column_widths(table, headers, rows)
                    for k, width in enumerate(column_widths):
                        for row in table.rows:
                            row.cells[k].width = width
                    
                    format_table(table)
                    doc.add_paragraph().paragraph_format.space_after = Pt(8)
            
            i = j - 1
        
        elif line and not line.startswith('<!--'):
            p = doc.add_paragraph()
            add_formatted_text(p, line)
            p.paragraph_format.space_after = Pt(8)
        
        i += 1
    
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    
    return docx_bytes.getvalue() 